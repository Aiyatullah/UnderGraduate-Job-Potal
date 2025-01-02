from flask import render_template, request, redirect, url_for, flash, session, jsonify, send_file
from . import app, mongo
from bson import ObjectId
from io import BytesIO
from docx import Document
import logging
import json
import os
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement

# Set up logging
logging.basicConfig(level=logging.INFO)

@app.route('/register', methods=['POST'])
def register():
    data = request.get_json()  # Get JSON data
    email = data['email']
    password = data['password']
    name = data['name']
    role = data['role']

    # Input validation
    if not email or not password or not name or not role:
        return jsonify({'message': 'All fields are required.'}), 400

    # Access the 'users' collection in the 'Hackathon' database
    users_collection = mongo.db.users

    # Check if the user already exists
    existing_user = users_collection.find_one({'email': email})
    if existing_user:
        return jsonify({'message': 'Email is already registered.'}), 400

    # Insert the new user into the 'users' collection
    new_user = {
        'email': email,
        'password': password,  # Store password directly (no hashing)
        'name': name,
        'role': role
    }
    
    try:
        users_collection.insert_one(new_user)
        return jsonify({'message': 'User registered successfully'}), 201
    except Exception as e:
        logging.error(f"Error inserting user: {e}")
        return jsonify({'message': 'An error occurred while registering the user.'}), 500

@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data['email']
    password = data['password']

    user = mongo.db.users.find_one({'email': email})
    if user and user['password'] == password:
        return jsonify({
            'message': 'Login successful',
            'user_id': str(user['_id']),
            'role': user['role']  # Include the role in the response
        }), 200
    else:
        return jsonify({'message': 'Login failed. Check your email and/or password'}), 401

@app.route('/profile/<user_id>', methods=['GET'])
def profile(user_id):
    try:
        user_id = ObjectId(user_id)
        user = mongo.db.users.find_one({'_id': user_id})
        if user:
            return jsonify({
                'name': user['name'],
                'email': user['email'],
                'role': user['role']  # Include the role in the response
            }), 200
        else:
            return jsonify({'message': 'User  not found'}), 404
    except Exception as e:
        return jsonify({'message': 'Invalid user ID'}), 400
    
@app.route('/logout', methods=['POST'])
def logout():
    # Clear the session data (log out the user)
    session.clear()
    return jsonify({'message': 'You have been logged out.'}), 200

from flask import Flask, request, jsonify
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement

@app.route('/api/upload-resume', methods=['POST'])
def upload_resume():
    try:
        # Parse incoming JSON data
        data = request.get_json()
        user_id = data.get('user_id')
        if not user_id:
            return jsonify({'error': 'user_id is required'}), 400
        try:
            user_id = ObjectId(user_id)
        except Exception as e:
            return jsonify({'error': 'Invalid user_id'}), 400


      # Extract fields and validate
        name = data.get('name', '').strip()
        skills = data.get('skills', '').strip()
        location = data.get('location', '').strip()
        degree = data.get('degree', '').strip()
        contact = data.get('contact', '').strip()
        email = data.get('email', '').strip()
        role = data.get('role', '').strip()
        description = data.get('description', '').strip()
        certifications = data.get('certifications', '').strip()
        languages = data.get('languages', '').strip()
        education = data.get('education', '').strip()
        experience = data.get('experience', '').strip()

        # Check for missing fields
        if not all([name, skills, location, degree, contact, email]):
            return jsonify({'error': 'Missing required fields'}), 400


        # Create a DOCX file in memory
        doc = Document()
        doc.add_heading('RESUME', level=1)

        # Set the font size for the entire document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(18)  # Set font size to 18

        # Name (bold, caps, blue color)
        paragraph = doc.add_paragraph()
        run_name = paragraph.add_run(f"{name.upper()}  ")  # Convert name to uppercase
        run_name.bold = True
        run_name.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
        run_email = paragraph.add_run(f"Email: {email}")
        run_email.italic = True
        run_email.font.color.rgb = RGBColor(0, 0, 255)  # Blue color

        # Contact and Location (same line)
        contact_location_paragraph = doc.add_paragraph()
        run_contact = contact_location_paragraph.add_run("Contact: ")
        run_contact.bold = True
        run_contact_location = contact_location_paragraph.add_run(f"{contact}  Location: {location}")
        run_contact_location.italic = True

        # Role
        role_paragraph = doc.add_paragraph()
        role_run = role_paragraph.add_run("ROLE: ")
        role_run.bold = True
        role_answer = role_paragraph.add_run(role)
        role_answer.italic = True

        # Description
        description_paragraph = doc.add_paragraph()
        desc_run = description_paragraph.add_run("Description: ")
        desc_run.bold = True
        desc_answer = description_paragraph.add_run(description)
        desc_answer.italic = True

        # Degree
        degree_paragraph = doc.add_paragraph()
        degree_run = degree_paragraph.add_run("Educational Degree: ")
        degree_run.bold = True
        degree_answer = degree_paragraph.add_run(degree)
        degree_answer.italic = True

        # Skills
        doc.add_paragraph("Skills:", style='Heading 2').bold = True
        for skill in skills.split(','):
            skill_paragraph = doc.add_paragraph(style='ListBullet')
            skill_run = skill_paragraph.add_run(skill.strip())
            skill_run.italic = True  # Skills in italics

        # Certifications
        doc.add_paragraph("Certifications:", style='Heading 2').bold = True
        for cert in certifications.split(','):
            cert_paragraph = doc.add_paragraph(style='ListBullet')
            cert_run = cert_paragraph.add_run(cert.strip())
            cert_run.italic = True  # Certifications in italics

        # Languages
        doc.add_paragraph("Languages:", style='Heading 2').bold = True
        for lang in languages.split(','):
            lang_paragraph = doc.add_paragraph(style='ListBullet')
            lang_run = lang_paragraph.add_run(lang.strip())
            lang_run.italic = True  # Languages in italics

        # Location (as a bullet point in bold and italic)
        doc.add_paragraph("Location:", style='Heading 2').bold = True
        location_paragraph = doc.add_paragraph()
        location_run = location_paragraph.add_run(location)
        location_run.bold = True
        location_run.italic = True

        # Save the document to a BytesIO object
        resume_stream = BytesIO()
        doc.save(resume_stream)
        resume_stream.seek(0) 

               # Update or insert the resume in MongoDB
        mongo.db.resumes.update_one(
            {'user_id': user_id},  # Query to find the existing resume
            {
                '$set': {  # Update these fields
                    'name': name,
                    'skills': skills,
                    'location': location,
                    'degree': degree,
                    'contact': contact,
                    'email': email,
                    'role': role,
                    'description': description,
                    'certifications': certifications,
                    'languages': languages,
                    'education': education,
                    'experience': experience,
                    'resume': resume_stream.getvalue()  # Store the binary content
                }
            },
            upsert=True  # Insert if the resume doesn't exist
        )

        return jsonify({'message': 'Resume generated and uploaded successfully!'}), 201
    except Exception as e:
        print("Error:", str(e))  # Log the error for debugging
        return jsonify({'error': 'Internal Server Error'}), 500
    

@app.route('/api/download-resume', methods=['GET'])
def download_resume():
    # Fetch the latest resume from MongoDB (you may want to modify this to fetch a specific user's resume)
    resume_entry = mongo.db.resumes.find_one(sort=[('_id', -1)])  # Get the most recent resume

    if resume_entry:
        resume_data = resume_entry['resume']
        resume_stream = BytesIO(resume_data)
        resume_stream.seek(0)

        # Send the file as a download
        return send_file(resume_stream, as_attachment=True, download_name='resume.doc', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        return jsonify({'message': 'No resume found.'}), 404

from flask import send_file, jsonify
from bson import ObjectId
from io import BytesIO

@app.route('/api/get-resume/<user_id>', methods=['GET'])
def get_resume(user_id):
    try:
        # Convert user_id to ObjectId
        user_id = ObjectId(user_id)
        
        # Fetch the user's resume from the resumes collection
        resume_entry = mongo.db.resumes.find_one({'user_id': user_id})
        
        if resume_entry:
            # Return the resume data in the expected format
            return jsonify({
                'skills': resume_entry.get('skills', ''),  # Default to empty string if not found
                'experience': resume_entry.get('experience', []),  # Default to empty list if not found
                'degree': resume_entry.get('degree', ''),  # Default to empty string if not found
                'certifications': resume_entry.get('certifications', ''),  # Default to empty string if not found
                'languages': resume_entry.get('languages', '')  # Default to empty string if not found
            }), 200
        else:
            return jsonify({'message': 'No resume found.'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500
